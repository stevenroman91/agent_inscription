"""
Module pour extraire le contenu des documents PDF et DOCX
"""
import os
from typing import List, Dict
from pathlib import Path
import pdfplumber
from docx import Document
from PIL import Image


class DocumentExtractor:
    """Extracteur de contenu pour différents types de documents"""
    
    def __init__(self, documents_dir: str = "."):
        self.documents_dir = Path(documents_dir)
    
    def extract_pdf(self, file_path: str) -> Dict[str, any]:
        """Extrait le texte d'un fichier PDF"""
        try:
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_content.append({
                            "page": page_num,
                            "text": text
                        })
            
            full_text = "\n\n".join([page["text"] for page in text_content])
            
            return {
                "type": "pdf",
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "pages": len(text_content),
                "text": full_text,
                "pages_content": text_content
            }
        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction du PDF {file_path}: {str(e)}")
    
    def extract_docx(self, file_path: str) -> Dict[str, any]:
        """Extrait le texte d'un fichier DOCX"""
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extraire aussi les tableaux
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append(table_data)
            
            full_text = "\n".join(paragraphs)
            
            return {
                "type": "docx",
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "text": full_text,
                "paragraphs": paragraphs,
                "tables": tables_content
            }
        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction du DOCX {file_path}: {str(e)}")
    
    def extract_image_info(self, file_path: str) -> Dict[str, any]:
        """Extrait les informations d'une image"""
        try:
            with Image.open(file_path) as img:
                return {
                    "type": "image",
                    "file_path": file_path,
                    "file_name": os.path.basename(file_path),
                    "format": img.format,
                    "size": img.size,
                    "mode": img.mode,
                    "width": img.width,
                    "height": img.height
                }
        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction de l'image {file_path}: {str(e)}")
    
    def extract_all_documents(self) -> List[Dict[str, any]]:
        """Extrait le contenu de tous les documents dans le répertoire"""
        documents = []
        
        for file_path in self.documents_dir.glob("*"):
            if file_path.is_file():
                try:
                    if file_path.suffix.lower() == ".pdf":
                        doc = self.extract_pdf(str(file_path))
                        documents.append(doc)
                    elif file_path.suffix.lower() == ".docx":
                        doc = self.extract_docx(str(file_path))
                        documents.append(doc)
                    elif file_path.suffix.lower() in [".jpg", ".jpeg", ".png"]:
                        doc = self.extract_image_info(str(file_path))
                        documents.append(doc)
                except Exception as e:
                    print(f"Erreur avec {file_path}: {str(e)}")
                    continue
        
        return documents

